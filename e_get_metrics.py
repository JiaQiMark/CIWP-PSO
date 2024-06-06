import cv2
import h5py
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
# from a_utils import used_tumer_list
from tqdm import tqdm
import matplotlib.pyplot as plt
import matplotlib
from e_quality_metrics import metric_functions


# 设置全局字体样式和大小
matplotlib.rcParams['font.family'] = 'serif'
matplotlib.rcParams['font.serif'] = ['Times New Roman']
matplotlib.rcParams['font.size'] = 10

# def psnr_calcu(original, compressed):
#     mse = np.mean((original - compressed) ** 2)
#     if mse == 0:
#         return float('inf')
#     n = 12
#     psnr_value = 20 * np.log10(2**n / np.sqrt(mse))
#     psnr_value = round(psnr_value, 4)
#     return psnr_value


def read_thresholds(exp_name, alg_name, img_i):
    th_list = []
    for nth in [4, 6, 8]:
        th_path = f"result//result_{alg_name}//{exp_name}//{alg_name}_{exp_name}-{img_i}_thresholds_{nth}.npy"
        th = np.load(th_path)
        th_list.append(th)
    return th_list


def read_tumer(img_i):
    matlab_data_path = f"D:\\000__program\\00__dataset\\01_img_Seg\\01__brainTumor_1_766\\{img_i}.mat"
    with h5py.File(matlab_data_path, 'r') as file:
        matlab_data = file['cjdata']
        label = np.array(matlab_data['label'])
        image = np.array(matlab_data['image'])
        patient_id = np.array(matlab_data['PID'])
        tumor_border = np.array(matlab_data['tumorBorder'])
        tumor_mask = np.array(matlab_data['tumorMask'])
        cvu8_image = cv2.convertScaleAbs(image)
    return image, label


def multi_threshold_segmentation(img, th):
    segmented = np.zeros_like(img)
    gray_value_max = np.amax(img)
    th.sort()

    m = len(th)
    for t in range(m + 1):
        start_t = 0 if t == 0 else th[t - 1]
        end_t = gray_value_max if t == m else th[t] - 1

        # 将灰度值在 [i * 10, (i+1) * 10 - 1] 范围内的像素点设置为对应的值
        segmented[(img >= start_t) & (img < end_t)] = start_t

    return segmented


def display_img(img_list):
    n = len(img_list)
    # 创建一个包含两个子图的图形
    fig, axles = plt.subplots(1, n, figsize=(18, 6))
    for i in range(n):
        axles[i].imshow(img_list[i], cmap='gray')
        # axles[i].set_title(f'T{thresholds.shape[0]}')
        # 调整子图之间的间距
    plt.tight_layout()
    plt.show()


def save_table(tumer_list, algn_list, metrics_v, docx_path, ths):
    doc = Document()
    # 1) 插入表头
    head_table = doc.add_table(rows=1, cols=len(algn_list) + 2)
    for col_num, header_text in enumerate(algn_list):
        head_table.cell(0, col_num + 2).text = header_text

    # 2) 逐个插入表格
    num_table = len(metrics_v)
    for i in range(num_table):
        img_metrics = metrics_v[i]
        img_metrics = img_metrics.T
        rows = img_metrics.shape[0]
        cols = img_metrics.shape[1]
        table = doc.add_table(rows, cols + 2)

        for row in range(rows):
            # 插入阈值
            cell_th = table.cell(row, 1)
            cell_th.text = str(ths[i][row])

            # 获取该行最大值的索引、最大值
            max_idx = np.argmax(img_metrics[row])
            if row == 0:
                cell_id = table.cell(row, 0)
                cell_id.text = str(tumer_list[i])
            for col in range(cols):
                cell = table.cell(row, col + 2)
                cell.text = str(img_metrics[row, col])
                cell.paragraphs[0].add_run('').font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if img_metrics[row, col] == img_metrics[row, max_idx]:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    cell.paragraphs[0].runs[0].font.bold = True

    doc.save(docx_path)


def g_psnr_ssim_fsim(metrics, tumer_list, docx_path):
    exp_list = ['bench', 'tumer']
    # algname_list = ['GWO', 'WOA', 'SSA', 'PSO', 'GA', 'FA', 'PSOsa', 'PSOssa', 'PSOmjq']
    algname_list = ['SSA', 'PSO', 'GA', 'PSOsa', 'PSOssa', 'PSOmjq']

    exp = exp_list[1]
    tumer_psnr = []
    thresholds = []
    for tumer_i in tqdm(tumer_list, 'Metrix docx generating...'):
        img_metrics = []
        for alg_i in range(len(algname_list)):
            alg = algname_list[alg_i]
            th_4, th_6, th_8 = read_thresholds(exp_name=exp, alg_name=alg, img_i=tumer_i)
            tumer, label = read_tumer(tumer_i)
            segmented_th4 = multi_threshold_segmentation(tumer, th_4)
            segmented_th6 = multi_threshold_segmentation(tumer, th_6)
            segmented_th8 = multi_threshold_segmentation(tumer, th_8)
            thresholds.append([th_4, th_6, th_8])
            # display_img([tumer, segmented_th4, segmented_th6, segmented_th8])

            metrics_th4 = metric_functions[metrics](tumer, segmented_th4)
            metrics_th6 = metric_functions[metrics](tumer, segmented_th6)
            metrics_th8 = metric_functions[metrics](tumer, segmented_th8)

            img_metrics.append([metrics_th4, metrics_th6, metrics_th8])
        tumer_psnr.append(np.array(img_metrics))

    # 将psrn保存为docx的表格
    save_table(tumer_list, algname_list, tumer_psnr, docx_path, thresholds)


if __name__ == '__main__':
    metrics_list = ['PSNR', 'SSIM', 'FSIM']
    metrics_idex = 2

    entropy_best = [1, 10, 17, 23, 44, 51, 54, 55, 60, 63, 66, 69, 70, 71, 72, 73, 80, 85, 86, 100,
                    117, 118, 121, 130, 133, 147, 151, 154, 165, 182, 186, 189, 192,
                    205, 215, 218, 226, 231, 232, 248, 267, 270, 279, 281, 286]

    id_psnr_ssim = [1, 60, 85, 182, 186, 248, 270, 281, 286]

    metrics = metrics_list[metrics_idex]
    docx_path = f"docx\\test_2_{metrics}.docx"
    g_psnr_ssim_fsim(metrics, id_psnr_ssim, docx_path)

    # used_tumer_list = [10, 17, 23, 44, 54, 69, 80, 85, 100]
    # tumer_list = []
    # for i in used_tumer_list:
    #     tumer, label = read_tumer(i)
    #     tumer_list.append(tumer)
    # display_img(tumer_list)

