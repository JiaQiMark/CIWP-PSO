import numpy as np
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


def algorithm_result(exp_name, alg_name, img_i):
    best_kapur = np.zeros(3)
    convergence_path = f"result//result_{alg_name}//{exp_name}//{alg_name}_{exp_name}-{img_i}_convergence_curve_4.npy"
    curve = np.load(convergence_path)
    best_kapur[0] = np.max(curve)

    convergence_path = f"result//result_{alg_name}//{exp_name}//{alg_name}_{exp_name}-{img_i}_convergence_curve_6.npy"
    curve = np.load(convergence_path)
    best_kapur[1] = np.max(curve)

    convergence_path = f"result//result_{alg_name}//{exp_name}//{alg_name}_{exp_name}-{img_i}_convergence_curve_8.npy"
    curve = np.load(convergence_path)
    best_kapur[2] = np.max(curve)
    best_kapur_4 = np.round(best_kapur, decimals=4)
    return best_kapur_4


if __name__ == '__main__':
    algname_list = ['GWO', 'WOA', 'SSA', 'PSO', 'GA', 'FA', 'PSOsa', 'PSOssa', 'PSOmjq']

    bench_or_tumer = 0

    bench_list = [1, 2, 5, 8, 9, 12, 16, 18, 21, 27, 28, 32, 34, 35, 37, 40]
    tumer_list = [1, 60, 85, 182, 186, 248, 270, 281, 286]
    img_list = [bench_list, tumer_list]

    exp_list = ['bench', 'tumer']
    exp = exp_list[bench_or_tumer]
    img_id = img_list[bench_or_tumer]

    docx_path = f"docx\\{exp}\\{exp}_kapur_4.docx"

    doc = Document()
    head_table = doc.add_table(rows=1, cols=len(algname_list) + 2)
    for col_num, header_text in enumerate(algname_list):
        head_table.cell(0, col_num + 2).text = header_text

    used_list = []
    for i in tqdm(img_id, 'docx generating...'):
        # 创建一个空的数组1
        empty_array = np.empty((3, 0))
        for an in range(len(algname_list)):
            algname = algname_list[an]
            kapur_entropy = algorithm_result(exp_list[bench_or_tumer], algname, i)
            empty_array = np.column_stack((empty_array, kapur_entropy))

        n_max = 0

        if empty_array[0][8] == np.max(empty_array[0]):
            n_max += 1
        if empty_array[1][8] == np.max(empty_array[1]):
            n_max += 1
        if empty_array[2][8] == np.max(empty_array[2]):
            n_max += 1

        # if n_max < 2:
        #     continue
        used_list.append(i)

        # 插入第一个表格
        rows, cols = empty_array.shape
        table1 = doc.add_table(rows, cols + 2)
        for row in range(rows):
            # 获取该行最大值的索引、最大值
            max_idx = np.argmax(empty_array[row])
            if row == 0:
                cell_id = table1.cell(row, 0)
                cell_id.text = str(i)
            for col in range(cols):
                cell = table1.cell(row, col + 2)
                cell.text = str(empty_array[row, col])
                cell.paragraphs[0].add_run('').font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if empty_array[row, col] == empty_array[row, max_idx]:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    cell.paragraphs[0].runs[0].font.bold = True

    # used_id_path = "docx\\used_tumer_id.npy"
    # np.save(used_id_path, np.array(used_list))
    print(used_list)
    doc.save(docx_path)




